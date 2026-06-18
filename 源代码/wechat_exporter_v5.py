"""
WeChat Chat Exporter - 微信聊天记录导出工具
Full-featured desktop GUI with WeChat-style HTML/Word/PDF export.
Copyright © 泪无痕
"""
import sys, os, sqlite3, json, csv, zstandard, re, html as h_mod, glob
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree as ET

OUT_DIR = r"D:\WeChatChatExporter\output"
dctx = zstandard.ZstdDecompressor()


def detect_wechat_paths():
    """Auto-detect all WeChat data paths on this machine."""
    results = []
    # Common WeChat data locations
    search_bases = []
    for drive in ['C', 'D', 'E', 'F']:
        for base in [
            f"{drive}:\\xwechat_files",
            f"{drive}:\\WeChat Files\\xwechat_files",
            os.path.expanduser(f"~\\xwechat_files"),
            os.path.expanduser(f"~\\Documents\\xwechat_files"),
            os.path.expanduser(f"~\\Desktop\\xwechat_files"),
        ]:
            if os.path.isdir(base) and base not in search_bases:
                search_bases.append(base)

    for base in search_bases:
        try:
            for entry in os.listdir(base):
                user_dir = os.path.join(base, entry)
                if not os.path.isdir(user_dir):
                    continue
                # Check for db_storage_decrypted or db_storage
                decrypted = os.path.join(user_dir, "db_storage_decrypted")
                raw = os.path.join(user_dir, "db_storage")
                contact_db = None
                if os.path.isdir(decrypted):
                    contact_db = os.path.join(decrypted, "contact", "contact.db")
                elif os.path.isdir(raw):
                    contact_db = os.path.join(raw, "contact", "contact.db")
                if contact_db and os.path.exists(contact_db):
                    # Read wxid from contact.db
                    wxid = entry.split('_')[0] if '_' in entry else entry
                    nickname = ""
                    try:
                        conn = sqlite3.connect(contact_db)
                        row = conn.execute(
                            "SELECT username, nick_name, remark FROM contact WHERE local_type=1 LIMIT 1"
                        ).fetchone()
                        if row:
                            wxid = row[0]
                            nickname = row[1] or row[2] or ""
                        conn.close()
                    except:
                        pass
                    db_path = decrypted if os.path.isdir(decrypted) else raw
                    label = f"{nickname} ({wxid})" if nickname else wxid
                    results.append({
                        'label': label,
                        'wxid': wxid,
                        'nickname': nickname,
                        'path': db_path,
                        'base': base,
                    })
        except:
            pass
    return results

def read_varint(data, pos):
    result = 0; shift = 0
    while pos < len(data):
        b = data[pos]; result |= (b & 0x7F) << shift; pos += 1
        if (b & 0x80) == 0: return result, pos
        shift += 7
    return result, pos

def is_readable(text):
    if not text: return False
    return sum(1 for c in text if c.isprintable() or c in '\n\r\t') / max(len(text), 1) > 0.5

def extract_text_from_protobuf(raw):
    if raw is None: return ""
    if isinstance(raw, (int, float)): return ""
    if not isinstance(raw, bytes): return str(raw)
    if raw[:4] != b'\x28\xb5\x2f\xfd':
        try: return raw.decode('utf-8')
        except: return ""
    try: decompressed = dctx.decompress(raw)
    except: return ""
    ci = decompressed.find(b':')
    rest = decompressed[ci+1:] if 0 < ci < 50 else decompressed
    texts = []; i = 0
    while i < len(rest):
        try:
            tag, ni = read_varint(rest, i)
            if ni <= i: i += 1; continue
            wire = tag & 0x07
            if wire == 2:
                length, ni2 = read_varint(rest, ni)
                if ni2 + length > len(rest): i = ni2 + 1; continue
                fd = rest[ni2:ni2+length]
                try:
                    t = fd.decode('utf-8')
                    if len(t) > 1 and is_readable(t): texts.append(t)
                except: pass
                i = ni2 + length
            elif wire == 0: _, i = read_varint(rest, ni)
            elif wire == 5: i = ni + 4
            elif wire == 1: i = ni + 8
            else: i = ni + 1
        except: i += 1
    return ' '.join(texts).strip()

def extract_wxid_from_content(content):
    if isinstance(content, bytes) and content[:4] == b'\x28\xb5\x2f\xfd':
        try:
            dec = dctx.decompress(content)
            ci = dec.find(b':')
            if 0 < ci < 50:
                return dec[:ci].decode('utf-8', errors='replace')
        except: pass
    return ""

def extract_wxid_from_xml(xml_text):
    from_wxid = ""; to_wxid = ""
    m1 = re.search(r'fromusername\s*=\s*"([^"]+)"', xml_text)
    m2 = re.search(r'tousername\s*=\s*"([^"]+)"', xml_text)
    if m1: from_wxid = m1.group(1)
    if m2: to_wxid = m2.group(1)
    return from_wxid, to_wxid

def decode_content(content, msg_type):
    if content is None: return ""
    if isinstance(content, (int, float)): return ""
    if not isinstance(content, bytes): return str(content)
    if content[:4] == b'\x28\xb5\x2f\xfd':
        try: decompressed = dctx.decompress(content)
        except: return "[解压失败]"
        try: text = decompressed.decode('utf-8', errors='replace')
        except: return ""
        if '<?xml' in text or '<msg' in text:
            return parse_xml_content(text, msg_type)
        result = extract_text_from_protobuf(content)
        return result
    try: return content.decode('utf-8')
    except: return ""

def parse_xml_content(xml_text, msg_type):
    xml_text = xml_text.strip()
    if not xml_text.startswith('<?xml'):
        xml_text = '<?xml version="1.0"?>' + xml_text
    try: root = ET.fromstring(xml_text)
    except:
        title = re.search(r'<title>(.*?)</title>', xml_text)
        if title: return title.group(1)
        return clean_xml_fallback(xml_text)
    msg = root if root.tag == 'msg' else root.find('.//msg')
    if msg is None: msg = root
    img = msg.find('.//img')
    if img is not None:
        w = img.get('cdnthumbwidth', '?')
        h = img.get('cdnthumbheight', '?')
        return f"[图片 {w}x{h}]"
    if msg.find('.//videomsg') is not None:
        v = msg.find('.//videomsg')
        dur = v.get('voicelength', '?')
        return f"[视频 {dur}ms]"
    voice = msg.find('.//voicemsg')
    if voice is not None:
        dur = voice.get('voicelength', '?')
        return f"[语音 {int(dur)//1000}秒]" if dur.isdigit() else "[语音]"
    emoji = msg.find('.//emoji')
    if emoji is not None:
        product = emoji.get('productid', '')
        if 'person' in product or 'sticker' in product.lower():
            return "[表情包]"
        return "[表情]"
    appmsg = msg.find('.//appmsg')
    if appmsg is not None:
        title_el = appmsg.find('title')
        title = title_el.text if title_el is not None and title_el.text else ""
        app_type = appmsg.find('type')
        type_val = app_type.text if app_type is not None else ""
        if type_val == '6':
            filename = appmsg.find('.//filename')
            if filename is not None and filename.text:
                return f"[文件] {filename.text}"
            return "[文件]"
        url_el = appmsg.find('url')
        url = url_el.text if url_el is not None else ""
        if title:
            if url and 'mp.weixin.qq.com' in url: return f"[公众号] {title}"
            elif url and 'weixin.qq.com' in url: return f"[小程序] {title}"
            elif url: return f"[链接] {title}"
            else: return f"[分享] {title}"
        return "[应用消息]"
    location = msg.find('.//location')
    if location is not None: return "[位置]"
    if msg.find('.//card') is not None: return "[名片]"
    if msg.find('.//transfer') is not None: return "[转账]"
    if msg.find('.//redenvelopes') is not None or msg.find('.//licious') is not None: return "[红包]"
    all_text = ''.join(root.itertext()).strip()
    if all_text and len(all_text) > 2:
        all_text = re.sub(r'\s+', ' ', all_text)
        if len(all_text) > 100: return all_text[:100] + "..."
        return all_text
    return ""

def clean_xml_fallback(xml_text):
    title = re.search(r'<title>(.*?)</title>', xml_text)
    if title: return title.group(1)
    if '<img' in xml_text: return "[图片]"
    if '<video' in xml_text: return "[视频]"
    if '<voice' in xml_text: return "[语音]"
    if '<emoji' in xml_text: return "[表情]"
    if '<location' in xml_text: return "[位置]"
    if '<appmsg' in xml_text: return "[应用消息]"
    return ""

MSG_TYPE_MAP = {
    1: 'text', 3: 'image', 34: 'voice', 42: 'card', 43: 'video',
    47: 'emoji', 48: 'location', 49: 'appmsg', 50: 'voip',
    10000: 'system', 10002: 'revoke',
}

MSG_TYPE_LABEL = {
    1: '文字', 3: '图片', 34: '语音', 42: '名片', 43: '视频',
    47: '表情', 48: '位置', 49: '应用消息', 50: '通话',
    10000: '系统消息', 10002: '撤回消息',
}


class WeChatDB:
    def __init__(self, db_dir):
        self.db_dir = db_dir
        self.contacts = {}
        self.self_wxid = ""
        self.conversations = {}
        self._load_contacts()
        self._load_sessions()
        self._load_messages()

    def _load_contacts(self):
        db = os.path.join(self.db_dir, "contact", "contact.db")
        if not os.path.exists(db): return
        conn = sqlite3.connect(db)
        for row in conn.execute("SELECT username, remark, nick_name, local_type FROM contact"):
            wxid, remark, nickname, ltype = row
            name = (remark.strip() if remark and remark.strip() else
                    (nickname.strip() if nickname and nickname.strip() else wxid))
            self.contacts[wxid] = name
            if ltype == 1 and not self.self_wxid: self.self_wxid = wxid
        conn.close()

    def _load_sessions(self):
        self.sessions = {}
        db = os.path.join(self.db_dir, "session", "session.db")
        if not os.path.exists(db): return
        conn = sqlite3.connect(db)
        try:
            for row in conn.execute("SELECT username, summary, last_msg_sender, last_sender_display_name FROM SessionTable"):
                self.sessions[row[0]] = {'summary': row[1], 'sender': row[2], 'sender_name': row[3]}
        except: pass
        conn.close()

    def _load_messages(self):
        db = os.path.join(self.db_dir, "message", "message_0.db")
        if not os.path.exists(db): return
        conn = sqlite3.connect(db)
        name2id = {}
        for row in conn.execute("SELECT rowid, user_name FROM Name2Id"):
            name2id[row[0]] = row[1]
        tables = conn.execute("SELECT name FROM sqlite_master WHERE type='table' AND name LIKE 'Msg_%'").fetchall()
        for (tname,) in tables:
            rows = conn.execute(f"""
                SELECT local_id, real_sender_id, create_time, message_content, local_type, source
                FROM [{tname}] ORDER BY sort_seq
            """).fetchall()
            messages = []
            wxid_participants = {}
            for row in rows:
                lid, sid, ct, content, mt, source = row
                dt = datetime.fromtimestamp(ct) if ct else datetime.now()
                text = decode_content(content, mt)
                sender_wxid = name2id.get(sid, "")
                if not sender_wxid:
                    sender_wxid = extract_wxid_from_content(content)
                if not sender_wxid and isinstance(content, bytes):
                    try:
                        if content[:4] == b'\x28\xb5\x2f\xfd':
                            dec = dctx.decompress(content)
                            txt = dec.decode('utf-8', errors='replace')
                            fw, tw = extract_wxid_from_xml(txt)
                            if fw: sender_wxid = fw
                    except: pass
                if text and ':' in text[:60]:
                    first_line_end = text.find('\n')
                    if first_line_end == -1: first_line_end = 60
                    prefix = text[:first_line_end]
                    if re.match(r'^wxid_\w+:', prefix) or re.match(r'^[a-z]\w{4,20}:', prefix):
                        text = text[first_line_end+1:].lstrip()
                if sender_wxid:
                    wxid_participants[sender_wxid] = wxid_participants.get(sender_wxid, 0) + 1
                sender_name = self.contacts.get(sender_wxid, "") if sender_wxid else ""
                if not sender_name and sender_wxid: sender_name = sender_wxid
                if not sender_name: sender_name = f"user_{sid}"
                messages.append({
                    'time': dt.strftime('%Y-%m-%d %H:%M:%S'),
                    'timestamp': ct,
                    'sender': sender_name,
                    'sender_wxid': sender_wxid,
                    'type': mt,
                    'type_name': MSG_TYPE_MAP.get(mt % 1000, 'unknown'),
                    'content': text,
                    'is_self': sender_wxid == self.self_wxid if sender_wxid else False,
                })
            if not messages: continue
            chat_name = self._resolve_chat_name(tname, wxid_participants)
            self.conversations[tname] = {
                'name': chat_name,
                'count': len(messages),
                'first': messages[0]['time'],
                'last': messages[-1]['time'],
                'messages': messages,
                'participants': wxid_participants,
            }
        conn.close()

    def _resolve_chat_name(self, tname, participants):
        others = {w: c for w, c in participants.items() if w != self.self_wxid}
        if len(others) == 0:
            if tname in self.sessions:
                s = self.sessions[tname]
                if s.get('sender_name'): return s['sender_name']
            return tname
        if len(others) == 1:
            wxid = list(others.keys())[0]
            return self.contacts.get(wxid, wxid)
        top = sorted(others.items(), key=lambda x: x[1], reverse=True)
        name = self.contacts.get(top[0][0], top[0][0])
        if len(others) > 2:
            name += f" 等{len(others)}人"
        return name

    def get_summary(self):
        return [{'table': t, 'name': v['name'], 'count': v['count'],
                 'first': v['first'], 'last': v['last']}
                for t, v in sorted(self.conversations.items(), key=lambda x: x[1]['count'], reverse=True)]


def filter_messages(messages, date_from, date_to):
    if not date_from and not date_to: return messages
    result = []
    for m in messages:
        ts = m['timestamp']
        if date_from and ts < date_from: continue
        if date_to and ts > date_to: continue
        result.append(m)
    return result


# ══════════════════════════════════════════
#  HTML Exporter (WeChat Style)
# ══════════════════════════════════════════
def export_html_wechat(messages, output_path, chat_name, self_name="我"):
    def esc(s):
        return h_mod.escape(str(s)) if s else ''
    def avatar(name, is_self):
        initial = name[0] if name else '?'
        colors = ['#07c160','#f5a623','#e74c3c','#3498db','#9b59b6','#1abc9c','#e67e22','#e91e63']
        c = '#07c160' if is_self else colors[sum(ord(ch) for ch in name) % len(colors)]
        return f'<div class="av" style="background:{c}">{esc(initial)}</div>'
    groups = []; cur = []; last_ts = 0
    for m in messages:
        if m['timestamp'] - last_ts > 300 and cur: groups.append(cur); cur = []
        cur.append(m); last_ts = m['timestamp']
    if cur: groups.append(cur)
    html = [f'''<!DOCTYPE html><html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1,user-scalable=no">
<title>{esc(chat_name)}</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
html,body{{height:100%;overflow:hidden}}
body{{background:#ededed;font-family:-apple-system,BlinkMacSystemFont,'PingFang SC','Microsoft YaHei',sans-serif;display:flex;justify-content:center}}
.phone{{width:100%;max-width:480px;background:#ededed;display:flex;flex-direction:column;height:100vh;height:100dvh;overflow:hidden;position:relative}}
@media(min-width:520px){{.phone{{border-radius:16px;box-shadow:0 8px 32px rgba(0,0,0,.12);margin:10px 0}}body{{background:#f0f0f0}}}}
.hdr{{background:linear-gradient(180deg,#303030,#3a3a3a);color:#fff;padding:env(safe-area-inset-top,0px) 16px 10px;display:flex;align-items:center;justify-content:center;position:relative;flex-shrink:0;min-height:48px}}
.hdr-back{{position:absolute;left:12px;font-size:20px;opacity:.7;padding:8px;cursor:pointer}}
.hdr-name{{font-size:17px;font-weight:600;max-width:60%;text-align:center;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
.hdr-dots{{position:absolute;right:12px;font-size:20px;opacity:.7;padding:8px;cursor:pointer}}
.chat{{flex:1;overflow-y:auto;-webkit-overflow-scrolling:touch;padding:12px 10px 20px;display:flex;flex-direction:column;gap:2px;padding-bottom:calc(20px + env(safe-area-inset-bottom,0px))}}
.tdiv{{text-align:center;margin:14px 0 6px}}
.tdiv span{{background:rgba(0,0,0,.06);color:#999;font-size:11px;padding:2px 8px;border-radius:4px;display:inline-block}}
.row{{display:flex;align-items:flex-start;gap:8px;padding:2px 4px}}
.row.s{{flex-direction:row-reverse}}
.av{{width:40px;height:40px;border-radius:6px;color:#fff;display:flex;align-items:center;justify-content:center;font-size:16px;font-weight:700;flex-shrink:0;-webkit-tap-highlight-color:transparent}}
.bub{{max-width:70%;padding:10px 13px;border-radius:6px;font-size:15px;line-height:1.5;word-break:break-all;white-space:pre-wrap;position:relative;margin-top:1px;word-wrap:break-word}}
.bub.o{{background:#fff;color:#1a1a1a}}
.bub.s{{background:#95ec69;color:#1a1a1a}}
.bub.o::before{{content:'';position:absolute;left:-5px;top:10px;border:5px solid transparent;border-right:5px solid #fff}}
.bub.s::after{{content:'';position:absolute;right:-5px;top:10px;border:5px solid transparent;border-left:5px solid #95ec69}}
.sn{{font-size:11px;color:#999;margin-bottom:2px}}
.sys{{text-align:center;margin:8px 0}}
.sys span{{background:rgba(0,0,0,.05);color:#999;font-size:11px;padding:2px 10px;border-radius:4px;display:inline-block}}
.img-box{{background:#f0f0f0;border-radius:6px;padding:16px;text-align:center;color:#999;font-size:13px;border:1px solid #e0e0e0;min-width:120px}}
.img-box .icon{{font-size:34px;margin-bottom:4px}}
.emoji-box{{font-size:30px;line-height:1.2}}
.file-box{{background:#f8f8f8;border:1px solid #e0e0e0;border-radius:6px;padding:10px 14px;display:flex;align-items:center;gap:10px}}
.file-box .icon{{font-size:30px}}
.file-box .info{{flex:1;min-width:0}}
.file-box .name{{font-size:13px;font-weight:500;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}}
.link-box{{background:#f8f8f8;border:1px solid #e0e0e0;border-radius:6px;padding:10px 14px}}
.link-box .title{{font-size:13px;font-weight:500;margin-bottom:4px}}
</style></head><body><div class="phone">
<div class="hdr"><div class="hdr-back">&#8249;</div><div class="hdr-name">{esc(chat_name)}</div><div class="hdr-dots">&#8943;</div></div>
<div class="chat">''']
    for group in groups:
        html.append(f'<div class="tdiv"><span>{esc(group[0]["time"])}</span></div>')
        seen_senders = set()
        for m in group:
            mt = m['type'] % 1000
            c = m['content']
            if mt in (10000, 10002):
                if c and '打招呼' not in c:
                    html.append(f'<div class="sys"><span>{esc(c)}</span></div>')
                continue
            role = 's' if m['is_self'] else 'o'
            nm = self_name if m['is_self'] else m['sender']
            html.append(f'<div class="row {role}">{avatar(nm, m["is_self"])}<div class="bub {role}">')
            if not m['is_self'] and m['sender_wxid'] not in seen_senders:
                html.append(f'<div class="sn">{esc(m["sender"])}</div>')
            seen_senders.add(m['sender_wxid'])
            if mt == 3:
                html.append('<div class="img-box"><div class="icon">🖼</div>图片</div>')
            elif mt == 43:
                html.append('<div class="img-box"><div class="icon">🎬</div>视频</div>')
            elif mt == 34:
                dur = ""
                dm = re.search(r'(\d+)', c) if c else None
                if dm: dur = f" {int(dm.group(1))//1000}秒"
                html.append(f'<div class="img-box"><div class="icon">🎤</div>语音{dur}</div>')
            elif mt == 47:
                html.append('<div class="emoji-box">[表情]</div>')
            elif mt == 49:
                if '[文件]' in c:
                    fname = c.replace('[文件]', '').strip()
                    html.append(f'<div class="file-box"><div class="icon">📄</div><div class="info"><div class="name">{esc(fname) if fname else "文件"}</div></div></div>')
                elif '[链接]' in c or '[公众号]' in c or '[小程序]' in c or '[分享]' in c:
                    html.append(f'<div class="link-box"><div class="title">{esc(c)}</div></div>')
                else:
                    html.append(f'<div class="img-box"><div class="icon">📎</div>{esc(c) if c else "应用消息"}</div>')
            elif mt == 48:
                html.append('<div class="img-box"><div class="icon">📍</div>位置</div>')
            elif mt == 42:
                html.append('<div class="img-box"><div class="icon">👤</div>名片</div>')
            elif mt == 50:
                html.append('<div class="img-box"><div class="icon">📞</div>通话</div>')
            elif c:
                if c.startswith('<?xml') or c.startswith('<'):
                    parsed = parse_xml_content(c, mt)
                    html.append(esc(parsed) if parsed else '<span style="color:#999">[消息]</span>')
                else:
                    html.append(esc(c))
            else:
                html.append(f'<span style="color:#999">[{m["type_name"]}]</span>')
            html.append('</div></div>')
    html.append('</div></div></body></html>')
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html))


# ══════════════════════════════════════════
#  Word Exporter (WeChat Bubble Style)
# ══════════════════════════════════════════
def export_word(messages, output_path, chat_name):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import copy

    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(chat_name)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(51, 51, 51)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"共 {len(messages)} 条消息")
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph("")

    COLORS = [
        (7, 193, 96), (245, 166, 35), (231, 76, 60), (52, 152, 219),
        (155, 89, 182), (26, 188, 156), (230, 126, 34), (233, 30, 99),
    ]

    def get_color(name):
        idx = sum(ord(c) for c in name) % len(COLORS)
        return COLORS[idx]

    def add_avatar_cell(table, row, col, name, is_self):
        cell = table.cell(row, col)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(name[0] if name else "?")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        if is_self:
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="07C160"/>')
        else:
            r, g, b = get_color(name)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{r:02X}{g:02X}{b:02X}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def set_cell_width(cell, width_cm):
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>')
        existing = tcPr.find(qn('w:tcW'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(tcW)

    def remove_cell_borders(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        existing = tcPr.find(qn('w:tcBorders'))
        if existing is not None:
            tcPr.remove(existing)
        tcPr.append(borders)

    last_date = ""
    for m in messages:
        mt = m['type'] % 1000
        dt_str = m['time']
        cur_date = dt_str[:10]

        # Date divider
        if cur_date != last_date:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(f"── {cur_date} ──")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(150, 150, 150)
            last_date = cur_date

        # System messages
        if mt in (10000, 10002):
            if m['content'] and '打招呼' not in m['content']:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(m['content'])
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(150, 150, 150)
                run.font.italic = True
            continue

        sender_name = m['sender']
        time_short = dt_str[11:16]
        is_self = m['is_self']

        # Content display
        if mt == 3:
            content_display = "[图片]"
        elif mt == 43:
            content_display = "[视频]"
        elif mt == 34:
            dur = ""
            dm = re.search(r'(\d+)', m['content']) if m['content'] else None
            if dm: dur = f" {int(dm.group(1))//1000}秒"
            content_display = f"[语音{dur}]"
        elif mt == 47:
            content_display = "[表情]"
        elif mt == 49:
            content_display = m['content'] if m['content'] else "[应用消息]"
        elif mt == 48:
            content_display = "[位置]"
        elif mt == 42:
            content_display = "[名片]"
        elif mt == 50:
            content_display = "[通话]"
        elif m['content']:
            content_display = m['content']
        else:
            content_display = f"[{MSG_TYPE_LABEL.get(mt, m['type_name'])}]"

        # Create bubble table: 3 columns (avatar | bubble | spacer for alignment)
        if is_self:
            # Right-aligned: spacer | bubble | avatar
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            # Remove borders from table
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tblBorders>'
            )
            existing = tblPr.find(qn('w:tblBorders'))
            if existing is not None:
                tblPr.remove(existing)
            tblPr.append(borders)

            spacer1 = table.cell(0, 0)
            bubble_cell = table.cell(0, 1)
            avatar_cell = table.cell(0, 2)

            set_cell_width(spacer1, 5)
            set_cell_width(bubble_cell, 8)
            set_cell_width(avatar_cell, 1.2)

            remove_cell_borders(spacer1)
            remove_cell_borders(bubble_cell)
            remove_cell_borders(avatar_cell)

            spacer1.text = ""

            # Avatar
            add_avatar_cell(table, 0, 2, sender_name, True)

            # Bubble content - green background
            set_cell_shading(bubble_cell, "95EC69")
            bp = bubble_cell.paragraphs[0]
            bp.paragraph_format.space_before = Pt(3)
            bp.paragraph_format.space_after = Pt(3)
            run = bp.add_run(content_display)
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'

        else:
            # Left-aligned: avatar | bubble | spacer
            table = doc.add_table(rows=1, cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tblBorders>'
            )
            existing = tblPr.find(qn('w:tblBorders'))
            if existing is not None:
                tblPr.remove(existing)
            tblPr.append(borders)

            avatar_cell = table.cell(0, 0)
            bubble_cell = table.cell(0, 1)
            spacer2 = table.cell(0, 2)

            set_cell_width(avatar_cell, 1.2)
            set_cell_width(bubble_cell, 8)
            set_cell_width(spacer2, 5)

            remove_cell_borders(avatar_cell)
            remove_cell_borders(bubble_cell)
            remove_cell_borders(spacer2)

            spacer2.text = ""

            # Avatar
            add_avatar_cell(table, 0, 0, sender_name, False)

            # Bubble content - white background
            set_cell_shading(bubble_cell, "FFFFFF")
            bp = bubble_cell.paragraphs[0]
            bp.paragraph_format.space_before = Pt(3)
            bp.paragraph_format.space_after = Pt(3)

            # Sender name in bubble (gray)
            snp = bubble_cell.add_paragraph()
            snp.paragraph_format.space_before = Pt(0)
            snp.paragraph_format.space_after = Pt(1)
            snr = snp.add_run(sender_name)
            snr.font.size = Pt(8)
            snr.font.color.rgb = RGBColor(150, 150, 150)

            run = bp.add_run(content_display)
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'

        # Small spacer after each message
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(1)
        sp.paragraph_format.space_after = Pt(1)

    doc.save(output_path)


# ══════════════════════════════════════════
#  PDF Exporter (WeChat Bubble Style)
# ══════════════════════════════════════════
def export_pdf(messages, output_path, chat_name):
    from fpdf import FPDF

    font_path = None
    for fp in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyh.ttf",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.exists(fp):
            font_path = fp
            break

    class WeChatPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=15)

        def header(self):
            pass

        def footer(self):
            self.set_y(-10)
            self.set_font('Chinese', '', 7) if font_path else self.set_font('Helvetica', '', 7)
            self.set_text_color(180, 180, 180)
            self.cell(0, 5, f'{self.page_no()}', align='C')

    pdf = WeChatPDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()

    if font_path:
        pdf.add_font('Chinese', '', font_path, uni=True)
        pdf.add_font('Chinese', 'B', font_path, uni=True)

    COLORS = [
        (7, 193, 96), (245, 166, 35), (231, 76, 60), (52, 152, 219),
        (155, 89, 182), (26, 188, 156), (230, 126, 34), (233, 30, 99),
    ]

    def get_color(name):
        idx = sum(ord(c) for c in name) % len(COLORS)
        return COLORS[idx]

    def font_name():
        return 'Chinese' if font_path else 'Helvetica'

    # Title bar (dark background)
    pdf.set_fill_color(48, 48, 48)
    pdf.rect(0, 0, 210, 22, 'F')
    pdf.set_font(font_name(), 'B', 14)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(0, 4)
    pdf.cell(210, 14, chat_name, align='C')

    pdf.set_xy(15, 28)

    page_w = 180
    avatar_size = 7
    bubble_max_w = 105
    bubble_pad = 3
    avatar_colors_self = (7, 193, 96)

    last_date = ""
    for m in messages:
        mt = m['type'] % 1000
        dt_str = m['time']
        cur_date = dt_str[:10]

        # Check page space
        if pdf.get_y() > 265:
            pdf.add_page()
            pdf.set_y(15)

        # Date divider
        if cur_date != last_date:
            pdf.set_font(font_name(), '', 8)
            pdf.set_text_color(150, 150, 150)
            pdf.set_x(60)
            pdf.cell(90, 8, f'── {cur_date} ──', align='C', ln=True)
            pdf.set_text_color(0, 0, 0)
            last_date = cur_date

        # System messages
        if mt in (10000, 10002):
            if m['content'] and '打招呼' not in m['content']:
                pdf.set_font(font_name(), '', 7)
                pdf.set_text_color(150, 150, 150)
                pdf.set_x(30)
                pdf.cell(150, 6, m['content'], align='C', ln=True)
                pdf.set_text_color(0, 0, 0)
            continue

        sender_name = m['sender']
        time_short = dt_str[11:16]
        is_self = m['is_self']

        # Content display
        if mt == 3:
            content_display = "[图片]"
        elif mt == 43:
            content_display = "[视频]"
        elif mt == 34:
            dur = ""
            dm = re.search(r'(\d+)', m['content']) if m['content'] else None
            if dm: dur = f" {int(dm.group(1))//1000}s"
            content_display = f"[语音{dur}]"
        elif mt == 47:
            content_display = "[表情]"
        elif mt == 49:
            content_display = m['content'] if m['content'] else "[应用消息]"
        elif mt == 48:
            content_display = "[位置]"
        elif mt == 42:
            content_display = "[名片]"
        elif mt == 50:
            content_display = "[通话]"
        elif m['content']:
            content_display = m['content']
        else:
            content_display = f"[{MSG_TYPE_LABEL.get(mt, m['type_name'])}]"

        # Measure text width
        pdf.set_font(font_name(), '', 9)
        text_w = pdf.get_string_width(content_display)
        bubble_w = min(text_w + bubble_pad * 2, bubble_max_w)
        row_h = max(avatar_size + 2, 8)

        # Truncate if needed
        if text_w > bubble_max_w - bubble_pad * 2:
            max_chars = int(len(content_display) * (bubble_max_w - bubble_pad * 2) / text_w)
            content_display = content_display[:max_chars] + "..."

        cur_y = pdf.get_y()

        if is_self:
            # Right-aligned bubble
            avatar_x = 15 + page_w - avatar_size
            bubble_x = avatar_x - bubble_w - 2

            # Avatar (green)
            pdf.set_fill_color(*avatar_colors_self)
            pdf.rect(bubble_x + bubble_w + 2, cur_y, avatar_size, avatar_size, 'F')
            pdf.set_font(font_name(), 'B', 7)
            pdf.set_text_color(255, 255, 255)
            pdf.set_xy(bubble_x + bubble_w + 2, cur_y)
            pdf.cell(avatar_size, avatar_size, sender_name[0] if sender_name else "?", align='C')

            # Bubble (green)
            pdf.set_fill_color(149, 236, 105)
            pdf.rect(bubble_x, cur_y, bubble_w, row_h, 'F')
            pdf.set_font(font_name(), '', 9)
            pdf.set_text_color(51, 51, 51)
            pdf.set_xy(bubble_x + bubble_pad, cur_y + 1)
            pdf.cell(bubble_w - bubble_pad * 2, row_h - 2, content_display)

        else:
            # Left-aligned bubble
            avatar_x = 15
            bubble_x = avatar_x + avatar_size + 2

            # Avatar (colored)
            r, g, b = get_color(sender_name)
            pdf.set_fill_color(r, g, b)
            pdf.rect(avatar_x, cur_y, avatar_size, avatar_size, 'F')
            pdf.set_font(font_name(), 'B', 7)
            pdf.set_text_color(255, 255, 255)
            pdf.set_xy(avatar_x, cur_y)
            pdf.cell(avatar_size, avatar_size, sender_name[0] if sender_name else "?", align='C')

            # Bubble (white)
            pdf.set_fill_color(255, 255, 255)
            pdf.rect(bubble_x, cur_y, bubble_w, row_h, 'F')

            # Sender name
            pdf.set_font(font_name(), '', 7)
            pdf.set_text_color(150, 150, 150)
            pdf.set_xy(bubble_x + bubble_pad, cur_y - 1)
            pdf.cell(bubble_w - bubble_pad * 2, 4, sender_name)

            # Content
            pdf.set_font(font_name(), '', 9)
            pdf.set_text_color(51, 51, 51)
            pdf.set_xy(bubble_x + bubble_pad, cur_y + 2)
            pdf.cell(bubble_w - bubble_pad * 2, row_h - 3, content_display)

        pdf.set_text_color(0, 0, 0)
        pdf.set_y(cur_y + row_h + 2)

    pdf.output(output_path)


def export_pdf_narrow(messages, output_path, chat_name):
    """手机友好的窄版PDF - 更窄更紧凑, 手机横屏/竖屏都好看"""
    from fpdf import FPDF

    font_path = None
    for fp in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\msyh.ttf",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if os.path.exists(fp):
            font_path = fp
            break

    class WeChatPDF(FPDF):
        def __init__(self):
            super().__init__()
            self.set_auto_page_break(auto=True, margin=10)
        def header(self):
            pass
        def footer(self):
            self.set_y(-8)
            self.set_font('Chinese', '', 7) if font_path else self.set_font('Helvetica', '', 7)
            self.set_text_color(180, 180, 180)
            self.cell(0, 4, f'{self.page_no()}', align='C')

    pdf = WeChatPDF()
    # 窄版: 左右边距18mm, 内容宽度只有 210-18*2=174mm
    pdf.set_margins(18, 10, 18)
    pdf.add_page()

    if font_path:
        pdf.add_font('Chinese', '', font_path, uni=True)
        pdf.add_font('Chinese', 'B', font_path, uni=True)

    COLORS = [
        (7, 193, 96), (245, 166, 35), (231, 76, 60), (52, 152, 219),
        (155, 89, 182), (26, 188, 156), (230, 126, 34), (233, 30, 99),
    ]

    def get_color(name):
        return COLORS[sum(ord(c) for c in name) % len(COLORS)]

    def fn():
        return 'Chinese' if font_path else 'Helvetica'

    # Title bar
    pdf.set_fill_color(48, 48, 48)
    pdf.rect(0, 0, 210, 18, 'F')
    pdf.set_font(fn(), 'B', 12)
    pdf.set_text_color(255, 255, 255)
    pdf.set_xy(0, 3)
    pdf.cell(210, 12, chat_name, align='C')

    pdf.set_xy(18, 22)

    # 窄版参数: 更紧凑但字体不缩太小
    avatar_size = 6
    bubble_max_w = 88
    bubble_pad = 3
    line_gap = 1.5
    msg_gap = 1.0

    last_date = ""
    for m in messages:
        mt = m['type'] % 1000
        dt_str = m['time']
        cur_date = dt_str[:10]

        if pdf.get_y() > 272:
            pdf.add_page()
            pdf.set_y(10)

        # Date divider
        if cur_date != last_date:
            pdf.set_font(fn(), '', 8)
            pdf.set_text_color(150, 150, 150)
            pdf.set_x(50)
            pdf.cell(110, 6, f'── {cur_date} ──', align='C', ln=True)
            pdf.set_text_color(0, 0, 0)
            last_date = cur_date

        # System messages
        if mt in (10000, 10002):
            if m['content'] and '打招呼' not in m['content']:
                pdf.set_font(fn(), '', 7)
                pdf.set_text_color(150, 150, 150)
                pdf.set_x(30)
                pdf.cell(150, 5, m['content'], align='C', ln=True)
                pdf.set_text_color(0, 0, 0)
            continue

        sender_name = m['sender']
        is_self = m['is_self']

        if mt == 3:
            content_display = "[图片]"
        elif mt == 43:
            content_display = "[视频]"
        elif mt == 34:
            dur = ""
            dm = re.search(r'(\d+)', m['content']) if m['content'] else None
            if dm: dur = f" {int(dm.group(1))//1000}s"
            content_display = f"[语音{dur}]"
        elif mt == 47:
            content_display = "[表情]"
        elif mt == 49:
            content_display = m['content'] if m['content'] else "[应用消息]"
        elif mt == 48:
            content_display = "[位置]"
        elif mt == 42:
            content_display = "[名片]"
        elif mt == 50:
            content_display = "[通话]"
        elif m['content']:
            content_display = m['content']
        else:
            content_display = f"[{MSG_TYPE_LABEL.get(mt, m['type_name'])}]"

        # Measure text width
        pdf.set_font(fn(), '', 9)
        text_w = pdf.get_string_width(content_display)
        bubble_w = min(text_w + bubble_pad * 2, bubble_max_w)
        row_h = max(avatar_size + 1, 7.5)

        # Truncate if needed
        if text_w > bubble_max_w - bubble_pad * 2:
            max_chars = int(len(content_display) * (bubble_max_w - bubble_pad * 2) / text_w)
            content_display = content_display[:max_chars] + "..."

        cur_y = pdf.get_y()
        content_w = 210 - 18 * 2  # 174mm

        if is_self:
            avatar_x = 18 + content_w - avatar_size
            bubble_x = avatar_x - bubble_w - 1.5

            pdf.set_fill_color(7, 193, 96)
            pdf.rect(bubble_x + bubble_w + 1.5, cur_y, avatar_size, avatar_size, 'F')
            pdf.set_font(fn(), 'B', 7)
            pdf.set_text_color(255, 255, 255)
            pdf.set_xy(bubble_x + bubble_w + 1.5, cur_y)
            pdf.cell(avatar_size, avatar_size, sender_name[0] if sender_name else "?", align='C')

            pdf.set_fill_color(149, 236, 105)
            pdf.rect(bubble_x, cur_y, bubble_w, row_h, 'F')
            pdf.set_font(fn(), '', 9)
            pdf.set_text_color(51, 51, 51)
            pdf.set_xy(bubble_x + bubble_pad, cur_y + 0.5)
            pdf.cell(bubble_w - bubble_pad * 2, row_h - 1, content_display)

        else:
            avatar_x = 18
            bubble_x = avatar_x + avatar_size + 1.5

            r, g, b = get_color(sender_name)
            pdf.set_fill_color(r, g, b)
            pdf.rect(avatar_x, cur_y, avatar_size, avatar_size, 'F')
            pdf.set_font(fn(), 'B', 7)
            pdf.set_text_color(255, 255, 255)
            pdf.set_xy(avatar_x, cur_y)
            pdf.cell(avatar_size, avatar_size, sender_name[0] if sender_name else "?", align='C')

            pdf.set_fill_color(255, 255, 255)
            pdf.rect(bubble_x, cur_y, bubble_w, row_h, 'F')

            pdf.set_font(fn(), '', 7)
            pdf.set_text_color(150, 150, 150)
            pdf.set_xy(bubble_x + bubble_pad, cur_y - 0.5)
            pdf.cell(bubble_w - bubble_pad * 2, 4, sender_name)

            pdf.set_font(fn(), '', 9)
            pdf.set_text_color(51, 51, 51)
            pdf.set_xy(bubble_x + bubble_pad, cur_y + 1.5)
            pdf.cell(bubble_w - bubble_pad * 2, row_h - 2, content_display)

        pdf.set_text_color(0, 0, 0)
        pdf.set_y(cur_y + row_h + msg_gap)

    pdf.output(output_path)


# ══════════════════════════════════════════
#  Standard Exporters
# ══════════════════════════════════════════
def export_csv(messages, path):
    with open(path, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.writer(f); w.writerow(['time','sender','type','content'])
        for m in messages: w.writerow([m['time'],m['sender'],m['type_name'],m['content']])

def export_json(messages, path):
    with open(path, 'w', encoding='utf-8') as f: json.dump(messages, f, ensure_ascii=False, indent=2, default=str)

def export_txt(messages, path):
    with open(path, 'w', encoding='utf-8') as f:
        for m in messages: f.write(f"[{m['time']}] {m['sender']}: {m['content']}\n\n")


# ══════════════════════════════════════════
#  Export Thread
# ══════════════════════════════════════════
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

class ExportThread(QThread):
    progress = pyqtSignal(str)
    done = pyqtSignal(str)
    def __init__(self, db, convs, out_dir, fmt, date_from, date_to):
        super().__init__()
        self.db = db; self.convs = convs; self.out_dir = out_dir
        self.fmt = fmt; self.date_from = date_from; self.date_to = date_to
    def run(self):
        os.makedirs(self.out_dir, exist_ok=True); ok = 0; fail = 0
        ext_map = {'html': 'html', 'word': 'docx', 'pdf': 'pdf', 'pdf_narrow': 'pdf', 'csv': 'csv', 'json': 'json', 'txt': 'txt'}
        for tname in self.convs:
            info = self.db.conversations[tname]
            msgs = filter_messages(info['messages'], self.date_from, self.date_to)
            name = info['name']
            safe = re.sub(r'[\\/:*?"<>|]', '_', name)
            base = os.path.join(self.out_dir, safe)
            ext = ext_map.get(self.fmt, self.fmt)
            self.progress.emit(f"导出: {name} ({len(msgs)}条)")
            try:
                if self.fmt == 'html': export_html_wechat(msgs, base + '.html', name)
                elif self.fmt == 'word': export_word(msgs, base + '.docx', name)
                elif self.fmt == 'pdf': export_pdf(msgs, base + '.pdf', name)
                elif self.fmt == 'pdf_narrow': export_pdf_narrow(msgs, base + '.pdf', name)
                elif self.fmt == 'csv': export_csv(msgs, base + '.csv')
                elif self.fmt == 'json': export_json(msgs, base + '.json')
                elif self.fmt == 'txt': export_txt(msgs, base + '.txt')
                ok += 1
            except Exception as e:
                self.progress.emit(f"  错误: {e}")
                fail += 1
        self.done.emit(f"完成! 成功 {ok} 个, 失败 {fail} 个\n输出目录: {self.out_dir}")


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("微信聊天记录导出工具 - Copyright © 泪无痕")
        self.setMinimumSize(1100, 700)
        self.resize(1200, 750)
        self.setStyleSheet("""
            QMainWindow{background:#f0f0f0}
            QGroupBox{font-weight:bold;border:1px solid #ddd;border-radius:8px;margin-top:12px;
                padding:16px 12px 12px;background:white}
            QGroupBox::title{subcontrol-origin:margin;left:16px;padding:0 6px;color:#07c160;font-size:13px}
            QListWidget{background:white;border:1px solid #e0e0e0;border-radius:4px;font-size:13px;padding:2px}
            QListWidget::item{padding:8px 10px;border-bottom:1px solid #f5f5f5}
            QListWidget::item:selected{background:#e6f9ee;color:#333}
            QListWidget::item:hover{background:#f8fdf9}
            QPushButton{background:#07c160;color:white;border:none;padding:8px 20px;border-radius:4px;
                font-size:13px;font-weight:bold}
            QPushButton:hover{background:#06ad56}
            QPushButton:pressed{background:#059a4c}
            QPushButton:disabled{background:#ccc}
            QLineEdit{border:1px solid #ddd;border-radius:4px;padding:6px 10px;font-size:13px;background:white}
            QComboBox{border:1px solid #ddd;border-radius:4px;padding:6px 10px;font-size:13px;background:white}
            QLabel{font-size:13px}
            QDateEdit{border:1px solid #ddd;border-radius:4px;padding:4px 8px;font-size:13px;background:white}
            QStatusBar{font-size:12px;color:#666}
        """)
        central = QWidget(); self.setCentralWidget(central)
        layout = QVBoxLayout(central); layout.setSpacing(6); layout.setContentsMargins(16, 10, 16, 10)

        hdr = QLabel("微信聊天记录导出工具 | Copyright © 泪无痕")
        hdr.setStyleSheet("color:#07c160;font-size:20px;font-weight:bold;padding:6px 0")
        hdr.setAlignment(Qt.AlignCenter); layout.addWidget(hdr)

        # Row 1: Account selector
        r1 = QHBoxLayout()
        r1.addWidget(QLabel("账号:"))
        self.accountCombo = QComboBox()
        self.accountCombo.setMinimumWidth(280)
        self.accountCombo.currentIndexChanged.connect(self.onAccountChanged)
        r1.addWidget(self.accountCombo, stretch=1)
        self.detectBtn = QPushButton("自动检测")
        self.detectBtn.setFixedWidth(100)
        self.detectBtn.setStyleSheet("background:#2196F3;font-size:12px;padding:4px 8px")
        self.detectBtn.clicked.connect(self.autoDetect)
        r1.addWidget(self.detectBtn)
        layout.addLayout(r1)

        # Row 2: Path + buttons
        r2 = QHBoxLayout()
        r2.setSpacing(6)
        self.pathEdit = QLineEdit()
        self.pathEdit.setPlaceholderText("数据库路径 (点击自动检测或手动浏览)")
        self.pathEdit.setMinimumWidth(400)
        r2.addWidget(self.pathEdit, stretch=3)
        b1 = QPushButton("浏览")
        b1.setFixedWidth(100)
        b1.setMinimumHeight(32)
        b1.setStyleSheet("font-size:13px;padding:6px 16px;background:#f5f5f5;border:1px solid #ccc;border-radius:4px")
        b1.clicked.connect(self.browse)
        r2.addWidget(b1)
        b2 = QPushButton("加载")
        b2.setFixedWidth(100)
        b2.setMinimumHeight(32)
        b2.setStyleSheet("font-size:13px;padding:6px 16px;background:#07C160;color:white;border:none;border-radius:4px")
        b2.clicked.connect(self.loadDB)
        r2.addWidget(b2)
        layout.addLayout(r2)

        body = QHBoxLayout()
        body.setSpacing(10)

        left = QGroupBox("会话列表")
        ll = QVBoxLayout(left)

        hb = QHBoxLayout()
        self.searchEdit = QLineEdit()
        self.searchEdit.setPlaceholderText("搜索联系人...")
        self.searchEdit.textChanged.connect(self.filterList)
        hb.addWidget(self.searchEdit)
        self.selAll = QPushButton("全选"); self.selAll.setFixedWidth(50)
        self.selAll.setStyleSheet("background:#f0f0f0;color:#333;border:1px solid #ddd;font-size:12px;padding:4px")
        self.selAll.clicked.connect(lambda: [self.convList.item(i).setSelected(True) for i in range(self.convList.count())])
        hb.addWidget(self.selAll)
        self.selNone = QPushButton("取消"); self.selNone.setFixedWidth(50)
        self.selNone.setStyleSheet("background:#f0f0f0;color:#333;border:1px solid #ddd;font-size:12px;padding:4px")
        self.selNone.clicked.connect(lambda: [self.convList.item(i).setSelected(False) for i in range(self.convList.count())])
        hb.addWidget(self.selNone)
        self.cntLabel = QLabel("0 个会话"); hb.addWidget(self.cntLabel)
        ll.addLayout(hb)

        self.convList = QListWidget(); self.convList.setSelectionMode(QAbstractItemView.MultiSelection)
        ll.addWidget(self.convList)
        body.addWidget(left, stretch=6)

        right = QGroupBox("导出设置")
        rl = QVBoxLayout(right)
        rl.setSpacing(8)

        rl.addWidget(QLabel("输出目录:"))
        self.outEdit = QLineEdit(OUT_DIR)
        rl.addWidget(self.outEdit)
        ob = QPushButton("选择目录...")
        ob.setFixedWidth(120)
        ob.setStyleSheet("font-size:12px;padding:4px 8px")
        ob.clicked.connect(self.browseOut)
        rl.addWidget(ob, alignment=Qt.AlignLeft)

        rl.addWidget(QLabel("导出格式:"))
        self.fmtCombo = QComboBox()
        self.fmtCombo.addItems(['html (微信风格)', 'word (.docx)', 'pdf', 'pdf (手机窄版)', 'csv', 'json', 'txt'])
        rl.addWidget(self.fmtCombo)

        rl.addWidget(QLabel("日期范围:"))
        dr = QHBoxLayout()
        dr.addWidget(QLabel("从:"))
        self.dateFrom = QDateEdit()
        self.dateFrom.setCalendarPopup(True)
        self.dateFrom.setDisplayFormat("yyyy-MM-dd")
        self.dateFrom.setDate(QDate(2020, 1, 1))
        self.dateFrom.setSpecialValueText("不限")
        dr.addWidget(self.dateFrom)
        dr.addWidget(QLabel("到:"))
        self.dateTo = QDateEdit()
        self.dateTo.setCalendarPopup(True)
        self.dateTo.setDisplayFormat("yyyy-MM-dd")
        self.dateTo.setDate(QDate.currentDate())
        self.dateTo.setSpecialValueText("不限")
        dr.addWidget(self.dateTo)
        rl.addLayout(dr)

        self.dateEnabled = QCheckBox("启用日期筛选")
        self.dateEnabled.setChecked(False)
        rl.addWidget(self.dateEnabled)

        rl.addStretch()
        self.exportBtn = QPushButton("开始导出"); self.exportBtn.setMinimumHeight(44)
        self.exportBtn.setStyleSheet("background:#07c160;font-size:15px;font-weight:bold;padding:8px")
        self.exportBtn.clicked.connect(self.startExport); rl.addWidget(self.exportBtn)

        self.progLabel = QLabel(""); self.progLabel.setWordWrap(True)
        self.progLabel.setStyleSheet("color:#666;margin-top:4px"); rl.addWidget(self.progLabel)
        self.statsLabel = QLabel(""); self.statsLabel.setStyleSheet("color:#999;font-size:11px")
        rl.addWidget(self.statsLabel)
        body.addWidget(right, stretch=4)
        layout.addLayout(body, stretch=1)

        footer = QLabel("仅用于个人数据备份，请勿用于非法用途 | 所有操作均在本地完成 | Copyright © 泪无痕")
        footer.setStyleSheet("color:#999;font-size:11px;padding:8px;text-align:center")
        footer.setAlignment(Qt.AlignCenter)
        layout.addWidget(footer)

        self.db = None; self.worker = None
        self.wechat_accounts = []

        QTimer.singleShot(500, self.autoDetect)

    def autoDetect(self):
        self.statusBar().showMessage("正在自动检测微信数据路径...")
        QApplication.processEvents()
        self.wechat_accounts = detect_wechat_paths()
        self.accountCombo.clear()
        if not self.wechat_accounts:
            self.accountCombo.addItem("未检测到微信数据")
            self.statusBar().showMessage("未检测到微信数据，请手动浏览选择路径")
            return
        for acc in self.wechat_accounts:
            self.accountCombo.addItem(acc['label'])
        self.statusBar().showMessage(f"检测到 {len(self.wechat_accounts)} 个微信账号")

    def onAccountChanged(self, idx):
        if 0 <= idx < len(self.wechat_accounts):
            self.pathEdit.setText(self.wechat_accounts[idx]['path'])

    def filterList(self, text):
        text = text.lower().strip()
        for i in range(self.convList.count()):
            item = self.convList.item(i)
            item.setHidden(text != "" and text not in item.text().lower())

    def browse(self):
        d = QFileDialog.getExistingDirectory(self, "选择数据库目录", self.pathEdit.text())
        if d: self.pathEdit.setText(d)
    def browseOut(self):
        d = QFileDialog.getExistingDirectory(self, "选择输出目录", self.outEdit.text())
        if d: self.outEdit.setText(d)

    def loadDB(self):
        db_dir = self.pathEdit.text()
        if not os.path.exists(db_dir):
            QMessageBox.warning(self, "错误", f"目录不存在:\n{db_dir}"); return
        self.statusBar().showMessage("正在加载数据库..."); QApplication.processEvents()
        self.db = WeChatDB(db_dir)
        self.convList.clear()
        for info in self.db.get_summary():
            self.convList.addItem(f"[{info['count']:>5}条] {info['name']}  ({info['first']} ~ {info['last']})")
        total = sum(s['count'] for s in self.db.get_summary())
        self.cntLabel.setText(f"{len(self.db.conversations)} 个会话, {total} 条消息")
        self.statsLabel.setText(f"自己的微信ID: {self.db.self_wxid}")
        self.statusBar().showMessage(f"加载完成: {len(self.db.conversations)} 个会话, {total} 条消息")

    def startExport(self):
        if not self.db:
            QMessageBox.warning(self, "错误", "请先加载数据库!"); return
        sel = self.convList.selectedItems()
        if not sel:
            QMessageBox.warning(self, "错误", "请选择要导出的会话!"); return
        convs = []
        all_t = [s['table'] for s in self.db.get_summary()]
        for item in sel:
            for t in all_t:
                if self.db.conversations[t]['name'] in item.text():
                    convs.append(t); break
        out_dir = self.outEdit.text()
        fmt_map = {'html (微信风格)': 'html', 'word (.docx)': 'word', 'pdf': 'pdf',
                   'pdf (手机窄版)': 'pdf_narrow', 'csv': 'csv', 'json': 'json', 'txt': 'txt'}
        fmt = fmt_map.get(self.fmtCombo.currentText(), 'html')

        date_from = None
        date_to = None
        if self.dateEnabled.isChecked():
            date_from = datetime(self.dateFrom.date().year(), self.dateFrom.date().month(), self.dateFrom.date().day()).timestamp()
            date_to = datetime(self.dateTo.date().year(), self.dateTo.date().month(), self.dateTo.date().day(), 23, 59, 59).timestamp()

        self.exportBtn.setEnabled(False)
        self.worker = ExportThread(self.db, convs, out_dir, fmt, date_from, date_to)
        self.worker.progress.connect(lambda m: self.progLabel.setText(m))
        self.worker.done.connect(self.onDone)
        self.worker.start()

    def onDone(self, msg):
        self.exportBtn.setEnabled(True); self.progLabel.setText(msg)
        self.statusBar().showMessage("导出完成!")
        QMessageBox.information(self, "导出完成", msg)


if __name__ == '__main__':
    app = QApplication(sys.argv); app.setStyle('Fusion')
    w = MainWindow(); w.show(); sys.exit(app.exec_())
