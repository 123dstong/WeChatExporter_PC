import os
import csv
import json
import re
from datetime import datetime
from typing import List, Dict, Optional


def format_timestamp(ts: int) -> str:
    """Convert Unix timestamp to readable string."""
    if not ts:
        return ""
    try:
        dt = datetime.fromtimestamp(ts)
        return dt.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        return str(ts)


def get_msg_type_name(msg_type: int, sub_type: int = 0) -> str:
    """Get human-readable message type name."""
    type_map = {
        1: "文本",
        3: "图片",
        34: "语音",
        42: "名片",
        43: "视频",
        47: "表情",
        48: "位置",
        49: "应用消息",
        50: "语音通话",
        51: "状态",
        52: "语音通话",
        53: "视频通话",
        62: "小视频",
        10000: "系统消息",
        10002: "系统消息",
    }
    return type_map.get(msg_type, f"类型{msg_type}")


def clean_content(content: str) -> str:
    """Clean message content for export."""
    if not content:
        return ""
    content = re.sub(r'<[^>]+>', '', content)
    content = re.sub(r'\[em_.*?\]', '[表情]', content)
    content = content.strip()
    return content


class ChatExporter:
    """Exports chat records to various formats."""

    def __init__(self, messages: List[Dict], contact_name: str,
                 output_dir: str, nickname_map: Optional[Dict] = None):
        self.messages = messages
        self.contact_name = contact_name
        self.output_dir = output_dir
        self.nickname_map = nickname_map or {}
        os.makedirs(output_dir, exist_ok=True)

    def _get_sender_name(self, msg: Dict) -> str:
        """Get sender display name."""
        if msg.get("is_sender"):
            return "我"
        talker = msg.get("talker", "")
        return self.nickname_map.get(talker, self.contact_name)

    def _safe_filename(self, name: str) -> str:
        """Create safe filename."""
        name = re.sub(r'[\\/:*?"<>|]', '_', name)
        name = name.strip('. ')
        return name[:100] if name else "chat"

    def export_csv(self) -> str:
        """Export to CSV format."""
        filename = f"{self._safe_filename(self.contact_name)}_聊天记录.csv"
        filepath = os.path.join(self.output_dir, filename)

        with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
            writer = csv.writer(f)
            writer.writerow(["时间", "发送者", "消息类型", "内容"])
            for msg in self.messages:
                writer.writerow([
                    format_timestamp(msg.get("time", 0)),
                    self._get_sender_name(msg),
                    get_msg_type_name(msg.get("type", 0), msg.get("sub_type", 0)),
                    clean_content(msg.get("content", "")),
                ])
        return filepath

    def export_html(self) -> str:
        """Export to HTML format with beautiful WeChat-like styling."""
        filename = f"{self._safe_filename(self.contact_name)}_聊天记录.html"
        filepath = os.path.join(self.output_dir, filename)

        html_parts = ["""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{contact} - 聊天记录</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'PingFang SC', 'Hiragino Sans GB', 'Microsoft YaHei', sans-serif;
    background: #f5f5f5;
    color: #333;
    line-height: 1.6;
}}
.header {{
    background: linear-gradient(135deg, #07c160, #06ad56);
    color: white;
    padding: 30px 20px;
    text-align: center;
    position: sticky;
    top: 0;
    z-index: 100;
    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
}}
.header h1 {{
    font-size: 22px;
    font-weight: 500;
    margin-bottom: 5px;
}}
.header .subtitle {{
    font-size: 13px;
    opacity: 0.85;
}}
.chat-container {{
    max-width: 800px;
    margin: 0 auto;
    padding: 20px 15px;
}}
.message-group {{
    margin-bottom: 2px;
}}
.time-divider {{
    text-align: center;
    padding: 15px 0 10px;
}}
.time-divider span {{
    background: rgba(0,0,0,0.06);
    color: #999;
    font-size: 12px;
    padding: 3px 12px;
    border-radius: 4px;
}}
.message {{
    display: flex;
    padding: 6px 15px;
    align-items: flex-start;
    gap: 10px;
}}
.message.sent {{
    flex-direction: row-reverse;
}}
.avatar {{
    width: 40px;
    height: 40px;
    border-radius: 6px;
    flex-shrink: 0;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 16px;
    font-weight: 500;
    color: white;
}}
.avatar.self {{
    background: linear-gradient(135deg, #07c160, #06ad56);
}}
.avatar.other {{
    background: linear-gradient(135deg, #5b9bd5, #4a8bc2);
}}
.bubble {{
    max-width: 65%;
    padding: 10px 14px;
    border-radius: 8px;
    position: relative;
    word-wrap: break-word;
    font-size: 14px;
    line-height: 1.6;
}}
.message.sent .bubble {{
    background: #95ec69;
    border-top-right-radius: 2px;
}}
.message.received .bubble {{
    background: white;
    border-top-left-radius: 2px;
    box-shadow: 0 1px 2px rgba(0,0,0,0.06);
}}
.bubble .sender {{
    font-size: 12px;
    color: #999;
    margin-bottom: 3px;
}}
.message.sent .bubble .sender {{
    text-align: right;
    color: rgba(0,0,0,0.35);
}}
.bubble .content {{
    word-break: break-word;
}}
.bubble .content.system {{
    text-align: center;
    color: #999;
    font-size: 12px;
    padding: 5px 10px;
    background: transparent;
    box-shadow: none;
}}
.bubble .msg-type {{
    display: inline-block;
    font-size: 12px;
    color: #999;
    border: 1px solid #ddd;
    padding: 2px 8px;
    border-radius: 4px;
    margin-top: 4px;
}}
.footer {{
    text-align: center;
    padding: 30px;
    color: #ccc;
    font-size: 12px;
}}
</style>
</head>
<body>
<div class="header">
    <h1>{contact}</h1>
    <div class="subtitle">聊天记录共 {count} 条</div>
</div>
<div class="chat-container">
""".format(contact=self.contact_name, count=len(self.messages))]

        last_time_str = ""
        for msg in self.messages:
            ts = msg.get("time", 0)
            time_str = format_timestamp(ts)
            current_date = time_str[:10] if time_str else ""

            if current_date and current_date != last_time_str[:10]:
                html_parts.append(
                    f'<div class="time-divider"><span>{current_date}</span></div>\n'
                )
                last_time_str = time_str

            is_sender = msg.get("is_sender", False)
            msg_type = msg.get("type", 1)
            content = clean_content(msg.get("content", ""))
            sender = self._get_sender_name(msg)
            css_class = "sent" if is_sender else "received"
            avatar_class = "self" if is_sender else "other"
            initial = sender[0] if sender else "?"

            if msg_type in (10000, 10002):
                html_parts.append(f'''
<div class="message {css_class}">
    <div class="bubble">
        <div class="content system">{content}</div>
    </div>
</div>
''')
            else:
                type_label = ""
                if msg_type not in (1,):
                    type_name = get_msg_type_name(msg_type)
                    type_label = f'<div class="msg-type">[{type_name}]</div>'

                html_parts.append(f'''
<div class="message {css_class}">
    <div class="avatar {avatar_class}">{initial}</div>
    <div class="bubble">
        <div class="sender">{sender}  <span style="font-size:11px;color:#bbb">{time_str}</span></div>
        <div class="content">{content if content else "&nbsp;"}</div>
        {type_label}
    </div>
</div>
''')

        html_parts.append("""
</div>
<div class="footer">由 WeChatChatExporter 导出</div>
</body>
</html>""")

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))
        return filepath

    def export_docx(self) -> str:
        """Export to Word DOCX format."""
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(10.5)

        title = doc.add_heading(level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"{self.contact_name} 聊天记录")
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(7, 193, 96)

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"共 {len(self.messages)} 条消息")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(153, 153, 153)

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ["时间", "发送者", "类型", "内容"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

        for msg in self.messages:
            row = table.add_row()
            row.cells[0].text = format_timestamp(msg.get("time", 0))
            row.cells[1].text = self._get_sender_name(msg)
            row.cells[2].text = get_msg_type_name(msg.get("type", 0))
            content = clean_content(msg.get("content", ""))
            row.cells[3].text = content[:500] if content else ""
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        filename = f"{self._safe_filename(self.contact_name)}_聊天记录.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def export_pdf(self) -> str:
        """Export to PDF format."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm, cm
        from reportlab.lib.colors import HexColor
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak, KeepTogether
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

        filename = f"{self._safe_filename(self.contact_name)}_聊天记录.pdf"
        filepath = os.path.join(self.output_dir, filename)

        font_registered = False
        font_names = [
            ('SimSun', r'C:\Windows\Fonts\simsun.ttc'),
            ('SimHei', r'C:\Windows\Fonts\simhei.ttf'),
            ('MSYaHei', r'C:\Windows\Fonts\msyh.ttc'),
            ('MSYaHeiBold', r'C:\Windows\Fonts\msyhbd.ttc'),
        ]
        for fname, fpath in font_names:
            if os.path.exists(fpath):
                try:
                    pdfmetrics.registerFont(TTFont(fname, fpath))
                    font_registered = True
                except Exception:
                    continue

        if font_registered:
            base_font = 'MSYaHei' if 'MSYaHei' in [f[0] for f in font_names] else 'Helvetica'
        else:
            base_font = 'Helvetica'

        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'ChatTitle',
            parent=styles['Heading1'],
            fontName=base_font,
            fontSize=20,
            textColor=HexColor('#07c160'),
            alignment=TA_CENTER,
            spaceAfter=6,
        )
        subtitle_style = ParagraphStyle(
            'ChatSubtitle',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=10,
            textColor=HexColor('#999999'),
            alignment=TA_CENTER,
            spaceAfter=20,
        )
        msg_self_style = ParagraphStyle(
            'MsgSelf',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=9,
            textColor=HexColor('#333333'),
            alignment=TA_RIGHT,
            spaceBefore=4,
            spaceAfter=4,
        )
        msg_other_style = ParagraphStyle(
            'MsgOther',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=9,
            textColor=HexColor('#333333'),
            alignment=TA_LEFT,
            spaceBefore=4,
            spaceAfter=4,
        )
        time_style = ParagraphStyle(
            'TimeDiv',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=8,
            textColor=HexColor('#999999'),
            alignment=TA_CENTER,
            spaceBefore=10,
            spaceAfter=6,
        )

        elements = []
        elements.append(Paragraph(f"{self.contact_name} - 聊天记录", title_style))
        elements.append(Paragraph(f"共 {len(self.messages)} 条消息 | 导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", subtitle_style))
        elements.append(Spacer(1, 10))

        green = HexColor('#07c160')
        light_gray = HexColor('#f5f5f5')
        last_date = ""

        for msg in self.messages:
            ts = msg.get("time", 0)
            time_str = format_timestamp(ts)
            current_date = time_str[:10] if time_str else ""

            if current_date and current_date != last_date:
                elements.append(Paragraph(f"—— {current_date} ——", time_style))
                last_date = current_date

            is_sender = msg.get("is_sender", False)
            sender = self._get_sender_name(msg)
            content = clean_content(msg.get("content", ""))
            msg_type = msg.get("type", 1)

            if msg_type in (10000, 10002):
                elements.append(Paragraph(
                    f'<font color="#999999">{content}</font>', time_style
                ))
                continue

            type_label = ""
            if msg_type not in (1,):
                type_name = get_msg_type_name(msg_type)
                type_label = f' <font color="#999999">[{type_name}]</font>'

            if is_sender:
                text = f'<b>我</b> <font size="7" color="#999999">{time_str}</font><br/>{content}{type_label}'
                elements.append(Paragraph(text, msg_self_style))
            else:
                text = f'<b>{sender}</b> <font size="7" color="#999999">{time_str}</font><br/>{content}{type_label}'
                elements.append(Paragraph(text, msg_other_style))

        elements.append(Spacer(1, 20))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontName=base_font,
            fontSize=8,
            textColor=HexColor('#cccccc'),
            alignment=TA_CENTER,
        )
        elements.append(Paragraph("由 WeChatChatExporter 导出", footer_style))

        doc.build(elements)
        return filepath

    def export_json(self) -> str:
        """Export to JSON format."""
        filename = f"{self._safe_filename(self.contact_name)}_聊天记录.json"
        filepath = os.path.join(self.output_dir, filename)

        data = {
            "contact": self.contact_name,
            "export_time": datetime.now().isoformat(),
            "message_count": len(self.messages),
            "messages": []
        }

        for msg in self.messages:
            data["messages"].append({
                "time": format_timestamp(msg.get("time", 0)),
                "timestamp": msg.get("time", 0),
                "sender": self._get_sender_name(msg),
                "is_sender": msg.get("is_sender", False),
                "type": msg.get("type", 0),
                "type_name": get_msg_type_name(msg.get("type", 0)),
                "content": clean_content(msg.get("content", "")),
            })

        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return filepath

    def export(self, format_type: str) -> str:
        """Export to specified format."""
        exporters = {
            "csv": self.export_csv,
            "html": self.export_html,
            "docx": self.export_docx,
            "pdf": self.export_pdf,
            "json": self.export_json,
        }
        exporter = exporters.get(format_type)
        if not exporter:
            raise ValueError(f"Unsupported format: {format_type}")
        return exporter()
